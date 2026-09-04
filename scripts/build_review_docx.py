#!/usr/bin/env python3
"""Fill the JCRT flyleaf-review template and append the review body.

usage: build_review_docx.py TEMPLATE.docx REVIEW.md OUT.docx
Front matter keys used: title, author, affiliation, doi, abstract, keywords.
Body: Markdown with ##/###/#### headings, *italics*, **bold**, [text](url)
links, and Pandoc-style footnotes ([^n] in the text, "[^n]: ..." definitions
anywhere on their own line).  Lines ending in two spaces keep a hard break.
The first "## " line is the bibliographic citation of the reviewed book.
"""
import re, sys
from pathlib import Path
import yaml
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

tpl, md_path, out = sys.argv[1:4]
raw = Path(md_path).read_text(encoding="utf-8")
_, fm, body = raw.split("---\n", 2)
meta = yaml.safe_load(fm)
FONT = "Book Antiqua"
LINK = RGBColor(0x00, 0x33, 0x66)
title = meta["title"]
author = meta["author"]
doi = meta.get("doi")
stable = f"https://doi.org/{doi}" if doi else meta.get("url", "")
running = f"{author.split()[-1]}: {title}"

# ---- footnote definitions ---------------------------------------------------
footnotes: dict[str, str] = {}
kept = []
for line in body.splitlines():
    m = re.match(r"^\[\^([^\]]+)\]:\s*(.*)$", line)
    if m:
        footnotes[m.group(1)] = m.group(2).strip()
    else:
        kept.append(line)
body = "\n".join(kept)
footnote_ids: dict[str, int] = {}


def set_font(run, size=None, bold=None, italic=None, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), name)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def replace_placeholder(doc, old, new):
    for p in doc.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                return True
    raise SystemExit(f"placeholder {old!r} not found")


def hyperlink(p, text, url, size=12, italic=False, part=None):
    part = part or p.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
    r = p.add_run(text); set_font(r, size, italic=italic)
    r.font.color.rgb = LINK; r.font.underline = True
    h.append(r._element); p._p.append(h)
    return r


def add_footnote_ref(p, key, size=12):
    """Superscript reference in the body + entry in footnotes.xml."""
    fid = footnote_ids.setdefault(key, len(footnote_ids) + 1)
    r = p.add_run(); set_font(r, size)
    r.font.superscript = True
    ref = OxmlElement("w:footnoteReference"); ref.set(qn("w:id"), str(fid))
    r._element.append(ref)


INLINE = r"(\*\*[^*]+\*\*|\*[^*]+\*|\[\^[^\]]+\]|\[[^\]]+\]\([^)]+\)|\n)"


def add_inline(p, text, size=12, base_bold=False, base_italic=False, part=None):
    """Emit runs for *italic* / **bold** / [^n] / [text](url) markdown."""
    for tok in re.split(INLINE, text):
        if not tok: continue
        if tok == "\n":
            p.add_run().add_break(WD_BREAK.LINE)
        elif tok.startswith("**"):
            set_font(p.add_run(tok[2:-2]), size, bold=True, italic=base_italic)
        elif tok.startswith("*"):
            set_font(p.add_run(tok[1:-1]), size, bold=base_bold, italic=not base_italic)
        elif tok.startswith("[^"):
            add_footnote_ref(p, tok[2:-1], size)
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            hyperlink(p, m.group(1), m.group(2), size, italic=base_italic, part=part)
        else:
            set_font(p.add_run(tok), size, bold=base_bold, italic=base_italic)


def add_field(run, instr):
    for tag, extra in (("w:fldChar", {"w:fldCharType": "begin"}), ("w:instrText", None), ("w:fldChar", {"w:fldCharType": "end"})):
        el = OxmlElement(tag)
        if extra:
            for k, v in extra.items(): el.set(qn(k), v)
        else:
            el.set(qn("xml:space"), "preserve"); el.text = instr
        run._element.append(el)


doc = Document(tpl)
replace_placeholder(doc, "[Review]", "Review")
replace_placeholder(doc, "[Title]", title)
replace_placeholder(doc, "[Author name(s)]", author)
replace_placeholder(doc, "[Permalink or DOI]", stable)
# make the stable URL a real hyperlink
for p in doc.paragraphs:
    if p.runs and p.runs[0].text.startswith("Stable URL"):
        r = p.runs[-1]
        part = doc.part
        rid = part.relate_to(stable, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
        r._element.addprevious(h); h.append(r._element)
        r.font.color.rgb = LINK; r.font.underline = True

# core properties
cp = doc.core_properties
cp.title = title; cp.author = author; cp.subject = meta.get("description", "")[:255]
cp.keywords = ", ".join(meta.get("keywords", [])); cp.category = "Book review"
cp.description = meta.get("abstract", "")[:255]; cp.language = "en-US"
cp.identifier = stable; cp.last_modified_by = "JCRT"

# ---- body section -------------------------------------------------------
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.header.is_linked_to_previous = False
sec.footer.is_linked_to_previous = False
sec.different_first_page_header_footer = False
sec.left_margin = sec.right_margin = Inches(1); sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run(running), 9, italic=True)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo = Path(tpl).parent / "x/word/media/image1.png"
if logo.exists():
    pic = fp.add_run().add_picture(str(logo), width=Inches(0.42))
    # alt text
    docpr = fp.runs[-1]._element.find(".//" + qn("wp:docPr"))
    docpr.set("descr", "a 2x2 grid with alternating black and red squares the letters JCRT. one letter per square in high contrast.")
    docpr.set("title", "JCRT logo")
set_font(fp.add_run("  Religious Theory"), 10, italic=True, name="Monotype Corsiva")
set_font(fp.add_run("  |  An Editorial Review Blog  |  Journal for Cultural and Religious Theory\t"), 8)
pr = fp.add_run(); set_font(pr, 8); add_field(pr, "PAGE")
# right tab stop for the page number
pPr = fp._element.get_or_add_pPr(); tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "9360"); tabs.append(tab); pPr.append(tabs)


def para(text="", align=None, size=12, space_after=6, indent=False, style=None):
    p = doc.add_paragraph(style=style)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if indent: p.paragraph_format.first_line_indent = Inches(0.5)
    if text: add_inline(p, text, size)
    return p

# opening block, as on previous JCRT reviews
p = para("BOOK REVIEW", WD_ALIGN_PARAGRAPH.CENTER, 12, 12, style="Heading 1")
lines = [l for l in body.strip().splitlines()]
i = 0
if lines and lines[0].startswith("## "):
    cite = lines[0][3:].strip(); i = 1
    cite = re.sub(r"^Book Review:\s*", "", cite, flags=re.I)
    p = para(None, WD_ALIGN_PARAGRAPH.CENTER, 12, 6)
    add_inline(p, cite.upper().replace("*", "*"), 12)   # italics preserved
para(f"Review by {author}", WD_ALIGN_PARAGRAPH.CENTER, 12, 18)

buf = []
def flush():
    global buf
    if buf:
        para(" ".join(buf).replace("\n ", "\n"), WD_ALIGN_PARAGRAPH.JUSTIFY, 12, 8, indent=True)
        buf = []
while i < len(lines):
    l = lines[i]; i += 1
    hard = l.endswith("  ")
    l = l.rstrip()
    if l.startswith("#### "):
        flush(); para(l[5:], None, 12, 6, style="Heading 3").runs[0].bold = True
    elif l.startswith("### "):
        flush(); para(l[4:], None, 12, 6, style="Heading 2").runs[0].bold = True
    elif not l.strip():
        flush()
    else:
        buf.append(l.strip() + ("\n" if hard else ""))
flush()

# heading styles: force serif font, black, sizes
for name, size in (("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)):
    st = doc.styles[name]; st.font.name = FONT; st.font.size = Pt(size); st.font.color.rgb = RGBColor(0, 0, 0)
    st.font.bold = True
    st.element.rPr.rFonts.set(qn("w:ascii"), FONT); st.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6); st.paragraph_format.keep_with_next = True
doc.styles["Heading 2"].paragraph_format.left_indent = Inches(0.5)
doc.styles["Heading 3"].paragraph_format.left_indent = Inches(0.5)

# signature block is the last body paragraph (name / affiliation / ORCID) -> left align, no indent, keep together
for p in doc.paragraphs[-1:]:
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.keep_together = True

# ---- footnotes part ---------------------------------------------------------
if footnote_ids:
    fn_part = next(pt for pt in doc.part.package.iter_parts() if str(pt.partname) == "/word/footnotes.xml")
    root = etree.fromstring(fn_part.blob)
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tmp = Document(tpl)  # scratch document only used to build runs
    for key, fid in sorted(footnote_ids.items(), key=lambda kv: kv[1]):
        fn = etree.SubElement(root, f"{{{W}}}footnote"); fn.set(f"{{{W}}}id", str(fid))
        p = tmp.add_paragraph()
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.0
        r = p.add_run(); set_font(r, 10); r.font.superscript = True
        r._element.append(OxmlElement("w:footnoteRef"))
        set_font(p.add_run(" "), 10)
        add_inline(p, footnotes.get(key, ""), 10, part=fn_part)
        fn.append(p._p)
    fn_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

doc.save(out)
print(out)
